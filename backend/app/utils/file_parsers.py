"""
File parsing utilities for extracting text from various file formats.
Supports PDF, DOCX, and TXT files.
"""
import io
from typing import Optional
import PyPDF2
import pdfplumber
from docx import Document


class FileParser:
    """Parser for extracting text from various file formats."""

    @staticmethod
    def parse_pdf(file_content: bytes) -> str:
        """
        Extract text from PDF file.

        Args:
            file_content: PDF file content as bytes

        Returns:
            Extracted text from PDF
        """
        text = ""

        # Try pdfplumber first (better for complex PDFs)
        try:
            with pdfplumber.open(io.BytesIO(file_content)) as pdf:
                for page in pdf.pages:
                    page_text = page.extract_text()
                    if page_text:
                        text += page_text + "\n"

            if text.strip():
                return text.strip()

        except Exception as e:
            print(f"pdfplumber failed: {e}, trying PyPDF2...")

        # Fallback to PyPDF2
        try:
            pdf_file = io.BytesIO(file_content)
            pdf_reader = PyPDF2.PdfReader(pdf_file)

            for page in pdf_reader.pages:
                page_text = page.extract_text()
                if page_text:
                    text += page_text + "\n"

            return text.strip()

        except Exception as e:
            print(f"PyPDF2 also failed: {e}")
            raise ValueError("Unable to extract text from PDF file")

    @staticmethod
    def parse_docx(file_content: bytes) -> str:
        """
        Extract text from DOCX file.

        Args:
            file_content: DOCX file content as bytes

        Returns:
            Extracted text from DOCX
        """
        try:
            doc = Document(io.BytesIO(file_content))
            text = []

            # Extract text from paragraphs
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text.append(paragraph.text)

            # Extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        if cell.text.strip():
                            text.append(cell.text)

            return "\n".join(text)

        except Exception as e:
            print(f"Error parsing DOCX: {e}")
            raise ValueError("Unable to extract text from DOCX file")

    @staticmethod
    def parse_txt(file_content: bytes) -> str:
        """
        Extract text from TXT file.

        Args:
            file_content: TXT file content as bytes

        Returns:
            Text content
        """
        try:
            # Try UTF-8 first
            return file_content.decode('utf-8')
        except UnicodeDecodeError:
            try:
                # Fallback to latin-1
                return file_content.decode('latin-1')
            except Exception as e:
                print(f"Error parsing TXT: {e}")
                raise ValueError("Unable to read text file")

    @staticmethod
    def parse_file(file_content: bytes, filename: str) -> str:
        """
        Parse file and extract text based on file extension.

        Args:
            file_content: File content as bytes
            filename: Original filename

        Returns:
            Extracted text

        Raises:
            ValueError: If file format is not supported or parsing fails
        """
        file_extension = filename.lower().split('.')[-1]

        if file_extension == 'pdf':
            return FileParser.parse_pdf(file_content)
        elif file_extension == 'docx':
            return FileParser.parse_docx(file_content)
        elif file_extension == 'txt':
            return FileParser.parse_txt(file_content)
        else:
            raise ValueError(f"Unsupported file format: {file_extension}. Supported formats: PDF, DOCX, TXT")


# Create singleton instance
file_parser = FileParser()
