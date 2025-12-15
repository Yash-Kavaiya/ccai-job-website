"""Document parsing service for resume text extraction."""

import io
import re
from typing import Dict, List, Optional, Tuple
from dataclasses import dataclass

import pdfplumber
from PyPDF2 import PdfReader
from docx import Document


@dataclass
class ParsedResume:
    """Parsed resume data structure."""
    raw_text: str
    sections: Dict[str, str]
    contact_info: Dict[str, str]
    skills: List[str]
    experience_entries: List[Dict]
    education_entries: List[Dict]
    certifications: List[str]
    languages: List[str]


class DocumentParser:
    """
    Document parser for extracting text and structured data from resumes.
    
    Supports PDF and DOCX formats.
    """
    
    # Common section headers
    SECTION_PATTERNS = {
        'experience': r'(?i)(work\s*experience|professional\s*experience|employment|work\s*history|experience)',
        'education': r'(?i)(education|academic|qualifications|degrees)',
        'skills': r'(?i)(skills|technical\s*skills|core\s*competencies|expertise|technologies)',
        'summary': r'(?i)(summary|profile|objective|about\s*me|professional\s*summary)',
        'certifications': r'(?i)(certifications?|certificates?|licenses?|credentials)',
        'projects': r'(?i)(projects?|portfolio|key\s*projects)',
        'languages': r'(?i)(languages?|language\s*skills)',
    }
    
    # Common tech skills for extraction
    TECH_SKILLS = [
        # Programming Languages
        'python', 'javascript', 'typescript', 'java', 'c++', 'c#', 'go', 'rust', 'ruby', 'php', 'swift', 'kotlin', 'scala', 'r',
        # AI/ML
        'machine learning', 'deep learning', 'nlp', 'natural language processing', 'computer vision', 'tensorflow', 'pytorch', 'keras', 'scikit-learn', 'pandas', 'numpy', 'opencv',
        # Cloud & DevOps
        'aws', 'azure', 'gcp', 'google cloud', 'docker', 'kubernetes', 'terraform', 'ansible', 'jenkins', 'ci/cd', 'devops',
        # Databases
        'sql', 'mysql', 'postgresql', 'mongodb', 'redis', 'elasticsearch', 'dynamodb', 'firestore', 'cassandra',
        # Web & Frameworks
        'react', 'angular', 'vue', 'node.js', 'express', 'django', 'flask', 'fastapi', 'spring', 'rails',
        # AI Platforms
        'openai', 'langchain', 'llm', 'gpt', 'bert', 'transformers', 'hugging face', 'dialogflow', 'amazon lex', 'ccai',
        # Data
        'data science', 'data engineering', 'etl', 'spark', 'hadoop', 'airflow', 'kafka', 'snowflake', 'databricks',
        # Other
        'git', 'linux', 'agile', 'scrum', 'rest api', 'graphql', 'microservices', 'api design',
    ]
    
    def parse_pdf(self, file_content: bytes) -> str:
        """Extract text from PDF using multiple methods for best results."""
        text_parts = []
        
        # Method 1: pdfplumber (better for complex layouts)
        try:
            with pdfplumber.open(io.BytesIO(file_content)) as pdf:
                for page in pdf.pages:
                    page_text = page.extract_text()
                    if page_text:
                        text_parts.append(page_text)
        except Exception:
            pass
        
        # Method 2: PyPDF2 (fallback)
        if not text_parts:
            try:
                reader = PdfReader(io.BytesIO(file_content))
                for page in reader.pages:
                    page_text = page.extract_text()
                    if page_text:
                        text_parts.append(page_text)
            except Exception:
                pass
        
        return '\n'.join(text_parts)
    
    def parse_docx(self, file_content: bytes) -> str:
        """Extract text from DOCX file."""
        try:
            doc = Document(io.BytesIO(file_content))
            paragraphs = [para.text for para in doc.paragraphs if para.text.strip()]
            
            # Also extract from tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        if cell.text.strip():
                            paragraphs.append(cell.text)
            
            return '\n'.join(paragraphs)
        except Exception:
            return ""
    
    def parse(self, file_content: bytes, filename: str) -> ParsedResume:
        """Parse resume file and extract structured data."""
        # Extract raw text based on file type
        if filename.lower().endswith('.pdf'):
            raw_text = self.parse_pdf(file_content)
        elif filename.lower().endswith(('.docx', '.doc')):
            raw_text = self.parse_docx(file_content)
        else:
            raw_text = file_content.decode('utf-8', errors='ignore')
        
        # Extract structured data
        sections = self._extract_sections(raw_text)
        contact_info = self._extract_contact_info(raw_text)
        skills = self._extract_skills(raw_text)
        experience = self._extract_experience(sections.get('experience', ''))
        education = self._extract_education(sections.get('education', ''))
        certifications = self._extract_certifications(sections.get('certifications', ''))
        languages = self._extract_languages(sections.get('languages', ''))
        
        return ParsedResume(
            raw_text=raw_text,
            sections=sections,
            contact_info=contact_info,
            skills=skills,
            experience_entries=experience,
            education_entries=education,
            certifications=certifications,
            languages=languages,
        )
    
    def _extract_sections(self, text: str) -> Dict[str, str]:
        """Extract resume sections based on headers."""
        sections = {}
        lines = text.split('\n')
        current_section = 'header'
        current_content = []
        
        for line in lines:
            line_stripped = line.strip()
            
            # Check if line is a section header
            found_section = None
            for section_name, pattern in self.SECTION_PATTERNS.items():
                if re.match(pattern, line_stripped) and len(line_stripped) < 50:
                    found_section = section_name
                    break
            
            if found_section:
                # Save previous section
                if current_content:
                    sections[current_section] = '\n'.join(current_content)
                current_section = found_section
                current_content = []
            else:
                current_content.append(line)
        
        # Save last section
        if current_content:
            sections[current_section] = '\n'.join(current_content)
        
        return sections
    
    def _extract_contact_info(self, text: str) -> Dict[str, str]:
        """Extract contact information from resume."""
        contact = {}
        
        # Email
        email_match = re.search(r'[\w\.-]+@[\w\.-]+\.\w+', text)
        if email_match:
            contact['email'] = email_match.group()
        
        # Phone
        phone_match = re.search(r'[\+]?[(]?[0-9]{1,3}[)]?[-\s\.]?[(]?[0-9]{1,4}[)]?[-\s\.]?[0-9]{1,4}[-\s\.]?[0-9]{1,9}', text)
        if phone_match:
            contact['phone'] = phone_match.group()
        
        # LinkedIn
        linkedin_match = re.search(r'linkedin\.com/in/[\w-]+', text, re.IGNORECASE)
        if linkedin_match:
            contact['linkedin'] = f"https://{linkedin_match.group()}"
        
        # GitHub
        github_match = re.search(r'github\.com/[\w-]+', text, re.IGNORECASE)
        if github_match:
            contact['github'] = f"https://{github_match.group()}"
        
        # Location (city, state pattern)
        location_match = re.search(r'([A-Z][a-z]+(?:\s[A-Z][a-z]+)?),?\s*([A-Z]{2})\b', text)
        if location_match:
            contact['location'] = f"{location_match.group(1)}, {location_match.group(2)}"
        
        return contact
    
    def _extract_skills(self, text: str) -> List[str]:
        """Extract technical skills from resume."""
        text_lower = text.lower()
        found_skills = []
        
        for skill in self.TECH_SKILLS:
            # Use word boundary matching
            pattern = r'\b' + re.escape(skill) + r'\b'
            if re.search(pattern, text_lower):
                found_skills.append(skill.title() if len(skill) > 3 else skill.upper())
        
        return list(set(found_skills))
    
    def _extract_experience(self, experience_text: str) -> List[Dict]:
        """Extract work experience entries."""
        entries = []
        
        # Pattern for company/role entries
        # Look for patterns like "Company Name | Role" or "Role at Company"
        lines = experience_text.split('\n')
        current_entry = None
        
        for line in lines:
            line = line.strip()
            if not line:
                continue
            
            # Check for date patterns (indicates new entry)
            date_match = re.search(r'((?:Jan|Feb|Mar|Apr|May|Jun|Jul|Aug|Sep|Oct|Nov|Dec)[a-z]*\.?\s*\d{4}|\d{4})\s*[-–—to]+\s*((?:Jan|Feb|Mar|Apr|May|Jun|Jul|Aug|Sep|Oct|Nov|Dec)[a-z]*\.?\s*\d{4}|\d{4}|Present|Current)', line, re.IGNORECASE)
            
            if date_match or (len(line) < 100 and any(c.isupper() for c in line[:20])):
                if current_entry and current_entry.get('title'):
                    entries.append(current_entry)
                
                current_entry = {
                    'title': line.split('|')[0].strip() if '|' in line else line,
                    'dates': date_match.group() if date_match else '',
                    'description': []
                }
            elif current_entry:
                current_entry['description'].append(line)
        
        if current_entry and current_entry.get('title'):
            entries.append(current_entry)
        
        # Convert description lists to strings
        for entry in entries:
            entry['description'] = ' '.join(entry['description'])
        
        return entries[:5]  # Limit to 5 most recent
    
    def _extract_education(self, education_text: str) -> List[Dict]:
        """Extract education entries."""
        entries = []
        
        # Common degree patterns
        degree_patterns = [
            r"(Bachelor'?s?|B\.?S\.?|B\.?A\.?|Master'?s?|M\.?S\.?|M\.?A\.?|Ph\.?D\.?|MBA|Associate'?s?)",
            r"(Computer Science|Engineering|Business|Mathematics|Physics|Data Science|Information Technology)"
        ]
        
        lines = education_text.split('\n')
        current_entry = {}
        
        for line in lines:
            line = line.strip()
            if not line:
                continue
            
            # Check for degree
            for pattern in degree_patterns:
                match = re.search(pattern, line, re.IGNORECASE)
                if match:
                    if current_entry:
                        entries.append(current_entry)
                    current_entry = {'degree': line, 'institution': '', 'year': ''}
                    break
            
            # Check for year
            year_match = re.search(r'\b(19|20)\d{2}\b', line)
            if year_match and current_entry:
                current_entry['year'] = year_match.group()
        
        if current_entry:
            entries.append(current_entry)
        
        return entries[:3]
    
    def _extract_certifications(self, cert_text: str) -> List[str]:
        """Extract certifications."""
        if not cert_text:
            return []
        
        certs = []
        lines = cert_text.split('\n')
        
        for line in lines:
            line = line.strip()
            if line and len(line) > 5 and len(line) < 150:
                # Clean up bullet points
                line = re.sub(r'^[\•\-\*\→]\s*', '', line)
                if line:
                    certs.append(line)
        
        return certs[:10]
    
    def _extract_languages(self, lang_text: str) -> List[str]:
        """Extract spoken languages."""
        common_languages = [
            'english', 'spanish', 'french', 'german', 'chinese', 'mandarin',
            'japanese', 'korean', 'portuguese', 'italian', 'russian', 'arabic',
            'hindi', 'dutch', 'swedish', 'polish'
        ]
        
        found = []
        text_lower = lang_text.lower() if lang_text else ''
        
        for lang in common_languages:
            if lang in text_lower:
                found.append(lang.title())
        
        return found
